"""
Genere un document .docx de preparation d'oral MSPR Bac+3 (frontend Health AI).
Sortie : C:\\Developper\\Ecole\\Health-AI-project\\preparation_oral_frontend.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_qa(doc, question, answer):
    p = doc.add_paragraph()
    r = p.add_run("Q : ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run(question)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("R : ")
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x27, 0x6F, 0x2A)
    p2.add_run(answer)
    doc.add_paragraph()


doc = Document()

# === PAGE DE GARDE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Health AI Platform")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Preparation orale - MSPR TPRE502 (Bac+3)")
rs.font.size = Pt(16)
rs.italic = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs2 = sub2.add_run("Volet Frontend - Next.js 16 / TypeScript / Shadcn UI")
rs2.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

intro_box = doc.add_paragraph()
intro_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
ri = intro_box.add_run(
    "Ce document recense l'ensemble du travail realise sur la partie frontend du projet "
    "Health AI : architecture, choix techniques, fonctionnalites, accessibilite, tests et livrables. "
    "Il est complete par une banque de questions / reponses pour anticiper l'oral."
)
ri.italic = True
ri.font.size = Pt(11)

doc.add_page_break()

# === SOMMAIRE ===
add_heading(doc, "Sommaire", 1)
sommaire = [
    "1. Contexte et objectifs du projet",
    "2. Architecture globale et place du frontend",
    "3. Stack technique frontend",
    "4. Pages developpees (parcours utilisateur)",
    "5. Systeme d'abonnements (Freemium / Premium / Premium+)",
    "6. Authentification et middleware",
    "7. Communication avec l'API (BFF)",
    "8. Composants metier (BMI, charts, suggestions)",
    "9. Accessibilite (WCAG 2.1 AA / RGAA)",
    "10. Tests E2E (Playwright + axe-core)",
    "11. Methodologie de developpement (Git, branches, PR)",
    "12. Documentation et livrables",
    "13. Limites et axes d'amelioration",
    "14. Banque de questions / reponses pour l'oral",
]
for s in sommaire:
    add_bullet(doc, s)

doc.add_page_break()

# === 1. CONTEXTE ===
add_heading(doc, "1. Contexte et objectifs du projet", 1)
add_para(doc,
    "Le projet Health AI Platform est un systeme de coaching sante propulse par l'IA. "
    "L'utilisateur peut suivre son alimentation (analyse de photos de repas), recevoir des plans "
    "de repas personnalises, des recommandations d'entrainement, et visualiser son evolution "
    "(poids, calories, macronutriments). Le projet adresse trois cibles : le grand public en B2C "
    "(freemium et premium) et les partenaires B2B (salles de sport, mutuelles, entreprises) en "
    "marque blanche."
)
add_para(doc,
    "Le cahier des charges (sujet MSPR TPRE502) impose un travail couvrant : conception, "
    "developpement, accessibilite (RGAA), tests, documentation et conduite du changement. "
    "Le frontend a ete realise avec Next.js 16 (React) et joue le role d'interface principale "
    "pour la version web."
)

add_heading(doc, "1.1 Objectifs frontend specifiques", 2)
for o in [
    "Offrir une experience fluide sur desktop et mobile (responsive).",
    "Respecter les criteres WCAG 2.1 AA / RGAA AA pour l'accessibilite.",
    "Differencier les fonctionnalites par tier (Free, Premium, Premium+).",
    "Supporter une marque blanche B2B via des themes dynamiques.",
    "Garantir un fonctionnement degrade lorsque le backend est indisponible (mode demo).",
    "Couvrir les parcours critiques par des tests automatises (Playwright).",
]:
    add_bullet(doc, o)

doc.add_page_break()

# === 2. ARCHITECTURE ===
add_heading(doc, "2. Architecture globale et place du frontend", 1)
add_para(doc,
    "L'architecture du projet est microservices. Elle separe les responsabilites en cinq services "
    "communiquant via HTTP/REST et gRPC :"
)
for a in [
    "health-next : application Next.js (web) - frontend principal.",
    "flutter-ai : application mobile Flutter (non couverte ici).",
    "backend-hono : BFF (Backend For Frontend) en Hono/TypeScript, port 3002. Orchestre les appels.",
    "engine-go : moteur metier en Go, expose des services gRPC (port 50051). Contient la logique IA et l'acces a PostgreSQL/pgvector.",
    "ia-python : service FastAPI pour la classification d'images alimentaires (PyTorch).",
]:
    add_bullet(doc, a)

add_para(doc,
    "Le frontend ne dialogue jamais directement avec engine-go ou ia-python : il passe systematiquement "
    "par le BFF backend-hono via des routes /api/*. Cela centralise l'authentification, le rate limiting "
    "et l'observabilite."
)

add_heading(doc, "2.1 Pourquoi un BFF ?", 2)
add_para(doc,
    "Le pattern Backend For Frontend permet d'avoir une API taillee pour les besoins de l'UI sans "
    "exposer la complexite des services internes. Concretement, le frontend fait un appel /api/home "
    "au BFF, qui agrege en parallele plusieurs appels gRPC vers engine-go (profil, stats, abonnement) "
    "et renvoie une reponse adaptee. Cela reduit le nombre de round-trips et evite de faire du "
    "fan-out cote client."
)

doc.add_page_break()

# === 3. STACK TECHNIQUE ===
add_heading(doc, "3. Stack technique frontend", 1)

stack = [
    ("Next.js 16", "Framework React moderne avec App Router, Server Components, Server Actions et streaming. Choisi pour le routing par fichiers, l'ecosysteme React, et la possibilite de rendu hybride (SSR + CSR)."),
    ("TypeScript", "Typage statique. Reduit les bugs en production, ameliore l'autocompletion et la maintenabilite."),
    ("Tailwind CSS 4", "Framework CSS utilitaire. Permet de styler directement dans le JSX, de garantir la coherence (couleurs, espacements) et d'utiliser des variables CSS pour le theming B2B."),
    ("Shadcn UI", "Bibliotheque de composants reutilisables construite au-dessus de Radix UI. Contrairement a une bibliotheque classique, le code des composants est copie dans le projet : on peut tout adapter."),
    ("Radix UI", "Primitives accessibles non stylees (Dialog, Select, Tabs, etc.). Garantit le respect des bonnes pratiques ARIA."),
    ("Zustand", "Gestion d'etat globale legere. Utilise pour le wizard d'inscription (persistance localStorage)."),
    ("React Hook Form + Zod", "Gestion des formulaires + validation par schema TypeScript. Validation cote client coherente avec le backend."),
    ("Recharts", "Bibliotheque de graphiques (LineChart, BarChart, PieChart). Utilisee pour les graphiques de poids, calories, macros."),
    ("Better Auth", "Solution d'authentification moderne, supporte les sessions, le hash des mots de passe et l'integration au middleware Next.js."),
    ("Lucide React", "Pack d'icones SVG legeres."),
    ("Playwright", "Framework de tests end-to-end. Simule un navigateur reel."),
    ("axe-core / @axe-core/playwright", "Audit d'accessibilite automatise."),
]
for name, desc in stack:
    p = doc.add_paragraph()
    rn = p.add_run(name + " : ")
    rn.bold = True
    p.add_run(desc)

doc.add_page_break()

# === 4. PAGES DEVELOPPEES ===
add_heading(doc, "4. Pages developpees (parcours utilisateur)", 1)

pages = [
    ("Landing page (/)", "Page d'accueil avec hero, presentation, et trois cartes d'abonnement (Free, Premium, Premium+). Animations hover (scale + shadow), responsive desktop/mobile.",
        ["Trois tiers d'abonnement avec prix (0EUR / 9.99EUR / 19.99EUR).",
         "Hover scale[1.02] + shadow pour les cartes, plus visible que hover:border-primary.",
         "CTA vers /inscription."]),
    ("Inscription (/inscription)", "Wizard a 5 etapes guidant l'utilisateur. Etat persiste dans localStorage via Zustand.",
        ["Etape 1 : age (slider).",
         "Etape 2 : poids et taille (FormField, validation Zod : poids 30-300, taille 100-250).",
         "Etape 3 : objectifs (perte/maintien/prise).",
         "Etape 4 : allergies (checkboxes, 'Aucune' auto-coche si rien d'autre).",
         "Etape 5 : signup (email + password) avec gestion explicite des erreurs (email deja utilise via 422).",
         "Recapitulatif (SummaryStep) avant validation finale."]),
    ("Connexion (/connexion)", "Page de login simple avec Better Auth. Redirige vers /dashboard.",
        ["Toaster pour feedback (succes/erreur).",
         "Liens vers /inscription si pas de compte."]),
    ("Dashboard principal (/dashboard)", "Page d'accueil apres connexion : 4 cards stats + carte BMI + 2 graphiques (poids, calories).",
        ["Stats : poids actuel, calories, proteines, activite.",
         "Carte BMI avec calcul automatique, categorie coloree, barre de progression et note sur les limites de l'IMC.",
         "Mode demo si l'API est down."]),
    ("Analytics (/dashboard/analytics)", "Vue analytique avancee : 4 cards resume, 3 onglets (Vue d'ensemble, Nutrition, Poids).",
        ["Onglet Nutrition : graphique macronutriments en donut.",
         "Onglet Poids : evolution sur 7/30/90 jours.",
         "Donnees demo si l'API est absente."]),
    ("Settings (/dashboard/settings)", "Reglages utilisateur : 3 onglets (Profil, Objectifs, Abonnement).",
        ["Profil : prenom, email, poids, taille.",
         "Objectifs : objectif calorique, repartition macros.",
         "Abonnement : tier actuel + bouton de surclassement."]),
    ("Nutrition - upload (/dashboard/nutrition)", "Upload d'une photo de repas, analyse via ia-python (FastAPI).",
        ["Affichage du resultat : aliment detecte + macros.",
         "Sauvegarde dans l'historique."]),
    ("Nutrition - historique (/dashboard/nutrition/history)", "Liste des repas analyses, filtres par periode.",
        ["Tabs : aujourd'hui, 7 jours, 30 jours, tout.",
         "Tri par date.",
         "6 repas demo en fallback."]),
    ("Nutrition - meal plan (/dashboard/nutrition/meal-plan)", "Plan de repas premium : 5 jours, repas par jour.",
        ["Onglets par jour, MealCard reutilisable.",
         "Protege par PremiumGuard (tier premium minimum)."]),
    ("Workouts (/dashboard/workouts)", "Programme d'entrainement Premium : 5 jours, exercices avec series/reps/RPE.",
        ["3 niveaux de difficulte.",
         "Connecte a POST /api/workout/generate.",
         "Protege par PremiumGuard."]),
    ("Clients B2B (/dashboard/clients)", "Page reservee Premium+ : 4 cards stats agregees (total, actifs 7j, taux Premium, calories moy.).",
        ["Acces filtre par usePremiumStatus().isPremiumPlus.",
         "Aucune liste nominative (donnees personnelles non exposees)."]),
]

for title_p, desc, bullets in pages:
    add_heading(doc, title_p, 2)
    add_para(doc, desc)
    for b in bullets:
        add_bullet(doc, b)
    doc.add_paragraph()

doc.add_page_break()

# === 5. TIERS ===
add_heading(doc, "5. Systeme d'abonnements (Freemium / Premium / Premium+)", 1)
add_para(doc,
    "Le projet definit trois niveaux d'abonnement avec des fonctionnalites differenciees, gerees "
    "uniformement cote frontend par un hook et un composant guard."
)

add_heading(doc, "5.1 Hook usePremiumStatus", 2)
add_para(doc,
    "Defini dans lib/hooks/use-premium-status.ts. Recupere le tier depuis l'API ou un fallback. "
    "Utilise un dictionnaire de niveaux numeriques (free=0, premium=1, premium_plus=2) pour gerer "
    "les comparaisons. Expose : tier, isLoading, isPremium, isPremiumPlus, et canAccess(requiredTier)."
)

add_heading(doc, "5.2 Composant PremiumGuard", 2)
add_para(doc,
    "Wrapper qui prend une prop requiredTier. Si l'utilisateur n'a pas le bon niveau, le contenu "
    "est rendu floute avec un overlay et un dialogue de surclassement (UpgradeDialog) qui adapte "
    "son message en fonction de l'ecart entre tier actuel et tier requis."
)

add_heading(doc, "5.3 Repartition des fonctionnalites", 2)
add_para(doc, "Free : dashboard de base, journal alimentaire, calcul d'IMC.", bold=True)
add_para(doc, "Premium : meal plan, recommandations workout, analytics avancees.", bold=True)
add_para(doc, "Premium+ (B2B) : page Clients (stats agregees), themes B2B en marque blanche.", bold=True)

doc.add_page_break()

# === 6. AUTH ===
add_heading(doc, "6. Authentification et middleware", 1)
add_para(doc,
    "L'authentification s'appuie sur Better Auth, une solution moderne pour Next.js. Lors de "
    "l'inscription, le frontend appelle /api/auth/sign-up (gere par Better Auth) puis /api/user/profile "
    "pour creer le profil sante. Une session est creee et stockee en cookie httpOnly."
)
add_para(doc,
    "Un middleware Next.js (middleware.ts) protege automatiquement les routes /dashboard/* : si le "
    "cookie de session est absent, l'utilisateur est redirige vers /connexion. Cela evite que des "
    "pages protegees ne s'affichent meme brievement avant le check cote client."
)

add_heading(doc, "6.1 Gestion des erreurs d'inscription", 2)
add_para(doc,
    "Cas particulier : email deja utilise. Le backend renvoie un 422 que la fonction apiFetch "
    "transforme en erreur typee. Le composant signup-step capte cette erreur et affiche un message "
    "explicite, evitant de rester bloque sur 'Inscription en cours' indefiniment."
)

doc.add_page_break()

# === 7. API ===
add_heading(doc, "7. Communication avec l'API (BFF)", 1)
add_para(doc,
    "Toutes les requetes vers le backend passent par lib/api.ts (fonction apiFetch). Cette fonction :"
)
for a in [
    "Ajoute automatiquement le prefixe /api et le base URL.",
    "Inclut les cookies de session (credentials: 'include').",
    "Transforme les reponses non-2xx en erreurs typees (status, required_tier, message).",
    "Gere un cache memoire pour les endpoints idempotents (utile pour limiter les appels repetes lors de la navigation).",
    "Permet aux composants d'avoir un fallback (donnees demo) si l'API est injoignable, sans crash.",
]:
    add_bullet(doc, a)

add_heading(doc, "7.1 Pattern mode demo", 2)
add_para(doc,
    "Chaque page critique (dashboard, analytics, settings, nutrition, workouts, clients) entoure son "
    "appel API d'un try/catch. En cas d'echec, elle bascule sur des donnees demo predeFINIES (DEMO_*). "
    "Cela permet de presenter le frontend meme sans backend en marche, et c'est tres utile pour la "
    "demo orale et les tests Playwright."
)

doc.add_page_break()

# === 8. COMPOSANTS METIER ===
add_heading(doc, "8. Composants metier", 1)

add_heading(doc, "8.1 BmiCard", 2)
add_para(doc,
    "Calcule l'IMC = poids / (taille en m)^2. Quatre categories : Insuffisance ponderale (<18.5), "
    "Normal (18.5-25), Surpoids (25-30), Obesite (>=30). Chaque categorie a sa couleur. Une barre "
    "de progression mappe l'IMC sur l'echelle 15-40. Une note avertit que l'IMC ne distingue pas "
    "muscle et graisse - un sportif peut avoir un IMC eleve sans exces de graisse."
)

add_heading(doc, "8.2 Graphiques (Recharts)", 2)
add_para(doc,
    "Trois graphiques principaux :"
)
for g in [
    "WeightEvolutionChart : LineChart, evolution du poids dans le temps.",
    "CaloriesChart : BarChart, calories par jour avec objectif.",
    "MacrosChart : PieChart en donut, repartition proteines / glucides / lipides.",
]:
    add_bullet(doc, g)

add_heading(doc, "8.3 MealSuggestions", 2)
add_para(doc,
    "Composant qui analyse les apports nutritionnels par rapport a des cibles et classe chaque "
    "nutriment en deficit (<80%), equilibre (80-120%) ou exces (>120%). Genere des suggestions "
    "personnalisees (par ex. 'Augmentez les proteines : oeufs, poulet, legumineuses')."
)

doc.add_page_break()

# === 9. ACCESSIBILITE ===
add_heading(doc, "9. Accessibilite (WCAG 2.1 AA / RGAA)", 1)
add_para(doc,
    "L'accessibilite est un critere du sujet MSPR. Le frontend respecte les niveaux WCAG 2.1 AA "
    "et RGAA AA pour permettre l'utilisation par des personnes en situation de handicap (visuel, "
    "moteur, cognitif)."
)

add_heading(doc, "9.1 Pratiques mises en place", 2)
for p in [
    "Hierarchie des titres : h1 unique par page, h2 pour les sections, structure logique.",
    "Headings sr-only (screen reader only) sur les sections muettes pour la navigation.",
    "aria-hidden='true' sur les icones decoratives (Lucide).",
    "Etiquettes explicites sur tous les inputs (label associe via htmlFor).",
    "focus-visible: au lieu de focus: pour ne pas afficher l'outline au clic souris.",
    "Contraste de couleur AA verifie dans toutes les combinaisons.",
    "Composants Radix : navigation clavier native (Tab, fleches, Echap, Entree).",
    "Tap targets >= 44x44 px sur mobile.",
    "alt sur toutes les images informatives.",
]:
    add_bullet(doc, p)

add_heading(doc, "9.2 Audit automatise", 2)
add_para(doc,
    "Tests Playwright + axe-core : un test (e2e/accessibility.spec.ts) parcourt 6 pages et verifie "
    "qu'aucune violation WCAG AA n'est detectee. Un script (scripts/generate-a11y-report.js) "
    "consolide les resultats en un rapport HTML avec correspondance WCAG/RGAA."
)

doc.add_page_break()

# === 10. TESTS ===
add_heading(doc, "10. Tests E2E (Playwright + axe-core)", 1)
add_para(doc,
    "Les tests end-to-end utilisent Playwright pour simuler un navigateur reel. Ils couvrent les "
    "parcours critiques de bout en bout : inscription, connexion, navigation dashboard, comportements "
    "freemium, accessibilite."
)

add_heading(doc, "10.1 Couverture", 2)
for t in [
    "auth-flow.spec.ts : inscription complete, connexion, deconnexion.",
    "dashboard-mocked.spec.ts : navigation entre les pages dashboard, mocks API.",
    "pages.spec.ts : verification de chaque page (titre, structure).",
    "user-journey.spec.ts : parcours utilisateur complet.",
    "accessibility.spec.ts : audit axe-core sur 6 pages.",
]:
    add_bullet(doc, t)

add_heading(doc, "10.2 Mocks", 2)
add_para(doc,
    "Pour ne pas dependre du backend en CI, certains tests interceptent les appels reseau "
    "(page.route(...)) et renvoient des reponses simulees. Cela rend les tests deterministes et rapides."
)

add_heading(doc, "10.3 CI/CD", 2)
add_para(doc,
    "Un workflow GitHub Actions (.github/workflows/ci.yml) execute les tests Playwright a chaque "
    "push sur main/develop et sur chaque pull request. Si un test casse, la PR ne peut pas etre "
    "fusionnee."
)

doc.add_page_break()

# === 11. METHODOLOGIE ===
add_heading(doc, "11. Methodologie de developpement", 1)
add_para(doc,
    "Le projet a ete developpe avec une approche feature-branch / pull request, une discipline "
    "courante en entreprise."
)

add_heading(doc, "11.1 Branches", 2)
for b in [
    "main : branche stable, recoit les merges via PR.",
    "frontend : branche d'integration des fonctionnalites frontend.",
    "Une branche par fonctionnalite : premium-plus, dashboard-analytics, calcul-imc, etc.",
    "Apres validation, merge dans frontend, puis frontend dans main.",
]:
    add_bullet(doc, b)

add_heading(doc, "11.2 Commit messages", 2)
add_para(doc,
    "Convention semi-conventionnelle : feat/fix/docs/test/chore + scope + description courte. "
    "Permet de generer un changelog et de comprendre l'historique sans lire le diff."
)

add_heading(doc, "11.3 Pull Requests", 2)
add_para(doc,
    "Chaque PR contient un descriptif minimaliste pour les collegues. Permet la revue de code et "
    "l'execution automatique des tests CI avant fusion."
)

doc.add_page_break()

# === 12. DOCUMENTATION ===
add_heading(doc, "12. Documentation et livrables", 1)

add_heading(doc, "12.1 Documents techniques", 2)
for d in [
    "docs/01-Migration-Shadcn-UI-DONE.md a docs/11-Page-Clients-B2B.md : 11 etapes documentees, chaque etape contient le pourquoi et le comment de chaque tache.",
    "docs/TODO.md : suivi global des taches, classees par priorite, avec tests manuels par tier.",
    "docs/livraison/01-Benchmark-Frontend.md : comparaison des frameworks etudies.",
    "docs/livraison/02-Documentation-Technique-Frontend.md : architecture detaillee.",
    "docs/livraison/03-Maquettes-Design-System.md : maquettes et design system.",
    "docs/livraison/04-Accessibilite-WCAG-RGAA.md : audit accessibilite.",
    "docs/livraison/05-Conduite-Changement.md : strategie de deploiement et adoption.",
    "docs/livraison/06-Rapport-Tests-Couverture.md : rapport de tests.",
]:
    add_bullet(doc, d)

add_heading(doc, "12.2 Livrables finaux", 2)
for d in [
    "PDF consolide : docs/livraison/HealthNext-Documentation-Frontend.pdf (toutes les sections en un seul document).",
    "documentation_frontend.docx : narration en style litteraire pour la livraison MSPR.",
    "23 screenshots dans docs/screenshots/ : illustrent chaque page et chaque etat de l'UI.",
    "Rapport d'accessibilite HTML genere automatiquement.",
]:
    add_bullet(doc, d)

doc.add_page_break()

# === 13. LIMITES ===
add_heading(doc, "13. Limites et axes d'amelioration", 1)
for l in [
    "Pas de role 'admin' implemente cote frontend : la page Clients est filtree par tier Premium+, mais aucun controle de role n'existe encore.",
    "Authentification : Better Auth est utilise pour la session, mais pas de 2FA.",
    "Pas d'internationalisation (i18n) : tout est en francais.",
    "Pas de PWA : pas de mode hors-ligne web (le mode hors-ligne est sur Flutter).",
    "Pas de tests unitaires (uniquement E2E). Une couche Vitest pourrait etre ajoutee pour les helpers et hooks.",
    "Le mode demo est utile mais peut masquer des bugs en l'absence de backend - il faut tester avec backend en preprod.",
    "Pas encore de monitoring frontend (Sentry, Datadog RUM).",
]:
    add_bullet(doc, l)

doc.add_page_break()

# === 14. QA ===
add_heading(doc, "14. Banque de questions / reponses pour l'oral", 1)
add_para(doc,
    "Les questions ci-dessous couvrent les sujets typiquement abordes en jury MSPR Bac+3 : "
    "techniques, methodologiques, business, accessibilite, securite et limites du projet."
)

# === 14.1 ARCHITECTURE ET STACK ===
add_heading(doc, "14.1 Architecture et choix techniques", 2)

add_qa(doc,
    "Pourquoi avoir choisi Next.js 16 plutot que React seul ?",
    "Next.js apporte le routing par fichiers (App Router), le rendu cote serveur (SSR), la generation "
    "statique (SSG), les Server Components et l'optimisation native des images et fonts. Avec React "
    "seul (Vite par exemple), il aurait fallu ajouter un router, configurer le SSR a la main, et "
    "perdre la fonction Server Action. Pour un projet de cette ampleur (10+ pages, SEO important), "
    "Next.js etait le bon compromis."
)

add_qa(doc,
    "Quelle est la difference entre Server Component et Client Component ?",
    "Un Server Component s'execute uniquement sur le serveur. Il peut acceder a la base de donnees "
    "directement, ne renvoie que du HTML au client, et ne peut pas utiliser useState ou useEffect. "
    "Un Client Component (marque 'use client') s'execute dans le navigateur, peut gerer l'etat et "
    "les evenements utilisateur. Sur ce projet, la majorite des pages dashboard sont des Client "
    "Components car elles doivent gerer l'etat de chargement et les fallbacks demo."
)

add_qa(doc,
    "Pourquoi Tailwind plutot que du CSS classique ou des CSS modules ?",
    "Tailwind reduit drastiquement le temps passe a nommer les classes et garantit une coherence "
    "visuelle (palette, spacings) car tout passe par le meme systeme de design tokens. Le bundle "
    "final reste petit grace au tree-shaking : seules les classes utilisees finissent dans le CSS. "
    "Avec du CSS classique, on aurait eu un fichier styles.css par composant et beaucoup plus de "
    "duplication."
)

add_qa(doc,
    "Pourquoi Shadcn UI et pas Material UI ou Ant Design ?",
    "Shadcn UI est une 'non-bibliotheque' : on copie le code des composants dans le projet plutot "
    "que de les installer comme dependance. Ca veut dire qu'on peut tout adapter sans hack et qu'on "
    "ne porte pas un poids enorme dans le bundle. C'est aussi base sur Radix UI, qui garantit "
    "l'accessibilite native (gestion clavier, ARIA). Material UI/Ant Design imposent leur design "
    "system et sont lourds (parfois 200+ Ko)."
)

add_qa(doc,
    "Qu'est-ce qu'un BFF et pourquoi en avoir un ?",
    "Backend For Frontend : c'est une couche d'API dediee aux besoins du client. Ici, le BFF "
    "(backend-hono) traduit les besoins de l'UI en appels gRPC vers engine-go. Avantages : on "
    "agrege plusieurs appels en un seul (moins de round-trips), on ajoute la logique d'auth/cache/"
    "rate limit a un seul endroit, et le frontend ne connait pas les details du moteur Go."
)

add_qa(doc,
    "Pourquoi gRPC entre BFF et engine-go au lieu de REST ?",
    "gRPC est plus rapide (HTTP/2, payloads binaires Protobuf), fortement type (les .proto generent "
    "des clients/serveurs typed dans plusieurs langages) et mieux adapte aux communications "
    "service-to-service. REST aurait eu plus d'overhead JSON et necessite une documentation OpenAPI "
    "manuelle pour rester en phase."
)

add_qa(doc,
    "Pourquoi Zustand plutot que Redux ou Context API ?",
    "Zustand est tres leger (1 Ko), n'a pas de boilerplate (pas de actions/reducers), supporte la "
    "persistance localStorage en une ligne. Redux aurait apporte une rigueur excessive pour notre "
    "besoin. Context API n'est pas concu pour de l'etat applicatif partage : il provoque des re-renders "
    "couteux. Pour le wizard d'inscription, Zustand etait l'outil le plus adapte."
)

# === 14.2 FONCTIONNALITES ===
add_heading(doc, "14.2 Fonctionnalites et logique metier", 2)

add_qa(doc,
    "Comment fonctionne le calcul de l'IMC ?",
    "IMC = poids (kg) / (taille (m))^2. Le poids et la taille sont collectes lors de l'inscription "
    "(wizard etape 2). Quatre categories : insuffisance ponderale (<18.5), normal (18.5-25), surpoids "
    "(25-30), obesite (>=30). Chaque categorie a sa couleur. J'ai ajoute une note importante : l'IMC "
    "ne distingue pas la masse musculaire de la masse grasse - un sportif peut avoir un IMC eleve "
    "sans exces de graisse."
)

add_qa(doc,
    "Comment se passe l'inscription utilisateur ?",
    "L'utilisateur passe par un wizard a 5 etapes (age, poids/taille, objectif, allergies, signup). "
    "L'etat du wizard est persiste dans localStorage via Zustand : si l'utilisateur quitte la page "
    "et revient, il retrouve sa progression. Au final, deux appels API : sign-up Better Auth pour "
    "creer le compte, puis POST /api/user/profile pour creer le profil sante."
)

add_qa(doc,
    "Pourquoi les 3 tiers d'abonnement ?",
    "Le cahier des charges impose le modele freemium pour le B2C, et un volet B2B en marque "
    "blanche. J'ai donc ajoute Premium+ qui adresse les partenaires B2B. Concretement : Free pour "
    "tester, Premium (9.99EUR) pour debloquer meal plan et workouts, Premium+ (19.99EUR) pour les "
    "salles de sport et entreprises avec acces a la page Clients."
)

add_qa(doc,
    "Comment proteger une page premium ?",
    "Deux niveaux : 1) Le hook usePremiumStatus expose canAccess(requiredTier) qu'on utilise dans "
    "la page elle-meme pour afficher un ecran 'Acces reserve' si le tier est insuffisant. 2) Le "
    "composant PremiumGuard qu'on utilise comme wrapper : il floute le contenu et ouvre un dialogue "
    "de surclassement. Le backend verifie aussi le tier sur les endpoints sensibles - c'est lui qui "
    "fait foi (defense en profondeur)."
)

add_qa(doc,
    "Comment gerez-vous les fonctionnalites B2B en marque blanche ?",
    "Quatre themes B2B sont definis dans le sidebar (Default, Sport, Mutuelle, Entreprise). Chaque "
    "theme modifie les variables CSS (couleur primaire, accent) au runtime via DynamicThemeProvider. "
    "C'est leger (juste du CSS variables) et ne necessite pas de rebuild."
)

add_qa(doc,
    "Pourquoi un mode demo si l'API ne repond pas ?",
    "Premier objectif : permettre la demo orale meme sans backend en marche. Deuxieme : permettre "
    "le travail front en parallele du back. Troisieme : Playwright peut tester sans dependre d'un "
    "service externe. Concretement, chaque page entoure son apiFetch d'un try/catch : si erreur, on "
    "bascule sur des donnees DEMO_* predefinies. Limite : il faut tester avec backend en preprod "
    "avant chaque livraison pour ne pas masquer un bug."
)

# === 14.3 ACCESSIBILITE ===
add_heading(doc, "14.3 Accessibilite", 2)

add_qa(doc,
    "Quelle est la difference entre WCAG et RGAA ?",
    "WCAG (Web Content Accessibility Guidelines) est le standard international du W3C, decline en "
    "trois niveaux A, AA, AAA. RGAA (Referentiel General d'Amelioration de l'Accessibilite) est le "
    "referentiel francais, base sur WCAG 2.1 AA et utilise par les administrations publiques. RGAA "
    "ajoute des criteres operationnels et un protocole de test plus strict. WCAG AA est aujourd'hui "
    "obligatoire pour les sites de service public et beaucoup d'entreprises privees."
)

add_qa(doc,
    "Comment testez-vous l'accessibilite ?",
    "Trois niveaux : 1) Audit manuel - navigation au clavier (Tab, Shift+Tab, Echap), zoom 200%, "
    "lecteur d'ecran (NVDA). 2) Audit automatise - axe-core via Playwright, qui verifie ~80% des "
    "criteres. 3) Verification visuelle - contraste de couleur (lighthouse, contrast checker). Les "
    "20% restants (lisibilite des labels, sens du contenu) ne peuvent etre audites qu'a la main."
)

add_qa(doc,
    "A quoi sert focus-visible: par rapport a focus: ?",
    "focus: applique le style des qu'un element est focus, y compris au clic souris. C'est moche "
    "et perturbant pour l'utilisateur souris. focus-visible: applique le style uniquement quand "
    "le focus arrive au clavier (Tab) ou par programmation. Resultat : les utilisateurs souris ne "
    "voient plus l'outline disgracieux au clic, mais les utilisateurs clavier gardent un signal "
    "visuel de leur position."
)

add_qa(doc,
    "Pourquoi mettre aria-hidden sur les icones ?",
    "Une icone Lucide est un SVG decoratif. Si on ne met pas aria-hidden='true', le lecteur d'ecran "
    "va annoncer 'image' ou tenter de lire le contenu, ce qui pollue. Avec aria-hidden='true', "
    "l'icone est ignoree par les technologies d'assistance. Le label texte adjacent porte le sens. "
    "Si l'icone est seule (par exemple un bouton fermer avec juste une croix), il faut a la place "
    "ajouter un aria-label='Fermer'."
)

add_qa(doc,
    "C'est quoi sr-only ?",
    "Une classe utilitaire Tailwind qui rend un texte invisible visuellement mais lu par les "
    "lecteurs d'ecran. Utilise pour les titres de section qui n'ont pas besoin d'etre vus mais "
    "structurent la page pour la navigation par titres. Exemple : <h2 className='sr-only'>Statistiques</h2>."
)

# === 14.4 TESTS ===
add_heading(doc, "14.4 Tests", 2)

add_qa(doc,
    "Quelle est la difference entre tests unitaires, d'integration et E2E ?",
    "Tests unitaires : verifient une fonction isolee (ex : un helper de calcul d'IMC). Tres rapides. "
    "Tests d'integration : verifient l'interaction entre plusieurs unites (ex : un composant React "
    "avec son hook). Tests E2E : simulent un utilisateur reel sur l'application complete (ex : "
    "remplir le wizard d'inscription jusqu'au bout). Plus lents mais plus realistes. Sur ce projet, "
    "j'ai privilegie les E2E avec Playwright car ils valident les parcours utilisateurs critiques."
)

add_qa(doc,
    "Pourquoi Playwright plutot que Cypress ou Selenium ?",
    "Playwright est plus rapide que Selenium (architecture moderne, sans WebDriver intermediaire), "
    "supporte plusieurs navigateurs (Chromium, Firefox, WebKit) avec une seule API, et a un meilleur "
    "support TypeScript. Cypress est limite a un seul navigateur a la fois et a des limites sur "
    "certaines interactions (frames cross-origin). Playwright est devenu le standard moderne."
)

add_qa(doc,
    "Comment mockez-vous l'API dans les tests ?",
    "page.route('**/api/**', route => route.fulfill({...})) : on intercepte toutes les requetes "
    "qui passent par /api et on renvoie une reponse simulee. Avantage : pas de dependance au "
    "backend, tests deterministes, executables en CI sans setup."
)

# === 14.5 SECURITE ===
add_heading(doc, "14.5 Securite", 2)

add_qa(doc,
    "Comment gerez-vous les sessions utilisateur ?",
    "Better Auth cree une session lors du sign-in/sign-up et stocke un cookie httpOnly + secure + "
    "sameSite. httpOnly empeche JavaScript d'y acceder (protection XSS). secure force HTTPS. "
    "sameSite='lax' protege contre CSRF en bloquant l'envoi automatique sur des requetes "
    "cross-site dangereuses."
)

add_qa(doc,
    "Comment vous protegez-vous du XSS ?",
    "React echappe par defaut tout le contenu inscrit avec {} dans le JSX. Si on injecte du HTML "
    "dynamique (dangerouslySetInnerHTML), il faut le sanitizer. Sur ce projet, on n'utilise pas "
    "dangerouslySetInnerHTML. Les contenus utilisateurs (nom, email) sont toujours interpoles via "
    "{} - donc echappes."
)

add_qa(doc,
    "Et contre le CSRF ?",
    "Cookie sameSite='lax' (envoye uniquement sur navigation top-level, jamais sur requetes "
    "automatiques cross-site). Better Auth ajoute aussi un token CSRF sur les operations sensibles. "
    "Le BFF verifie l'origine des requetes."
)

add_qa(doc,
    "A quoi sert le middleware Next.js ?",
    "Il s'execute avant chaque requete sur les routes correspondantes. Ici, middleware.ts protege "
    "/dashboard/* : si pas de cookie de session, on redirige vers /connexion. Avantage par rapport "
    "a un check cote client : la page protegee ne s'affiche jamais, meme brievement, ce qui evite "
    "les fuites de contenu sensible."
)

# === 14.6 PERFORMANCE ===
add_heading(doc, "14.6 Performance et optimisation", 2)

add_qa(doc,
    "Comment optimisez-vous le chargement de la page ?",
    "Plusieurs leviers : 1) Code splitting automatique de Next.js : chaque page est un bundle "
    "separe. 2) Optimisation des images via next/image (resize, formats modernes WebP/AVIF). 3) "
    "Cache memoire dans apiFetch pour les endpoints idempotents. 4) Tailwind purge les classes "
    "non utilisees, le CSS final fait quelques Ko. 5) Police optimisee via next/font."
)

add_qa(doc,
    "Comment monitoreriez-vous l'app en prod ?",
    "Plusieurs outils complementaires : 1) Sentry pour les erreurs JS et les erreurs reseau. 2) "
    "Vercel Analytics ou Datadog RUM pour les Core Web Vitals (LCP, FID, CLS). 3) Lighthouse CI "
    "dans la pipeline pour eviter les regressions de performance. Aujourd'hui c'est un axe "
    "d'amelioration : pas de monitoring en place."
)

# === 14.7 METHODOLOGIE ===
add_heading(doc, "14.7 Methodologie et organisation", 2)

add_qa(doc,
    "Comment avez-vous organise votre travail ?",
    "Workflow feature-branch : une branche par fonctionnalite (premium-plus, dashboard-analytics, "
    "calcul-imc...), commits semi-conventionnels, merge dans frontend apres validation, puis frontend "
    "dans main. J'ai aussi tenu une TODO liste (docs/TODO.md) avec les taches priorisees et les "
    "tests manuels par tier."
)

add_qa(doc,
    "Pourquoi documenter chaque etape ?",
    "1) Etre capable d'expliquer mes choix au jury et a un futur mainteneur. 2) Permettre la revue "
    "de code sans avoir a decoder le diff. 3) Sujet MSPR exige des livrables documentaires. Chaque "
    "etape (docs/01 a docs/11) explique le pourquoi de la tache et le comment de la realisation."
)

add_qa(doc,
    "Si vous deviez recommencer, que changeriez-vous ?",
    "1) Ajouter des tests unitaires Vitest sur les helpers (BMI, analyzeNutrients, getBmiCategory). "
    "2) Implementer un vrai systeme de roles (admin / partenaire B2B) plutot que de juste filtrer "
    "par tier. 3) Ajouter Sentry des le debut. 4) Mettre en place i18n pour preparer une expansion "
    "multi-langue. 5) Faire des Server Components pour les pages purement affichage (gain de "
    "performance et de SEO)."
)

# === 14.8 BUSINESS ===
add_heading(doc, "14.8 Business model et conduite du changement", 2)

add_qa(doc,
    "Quel est le modele economique du projet ?",
    "Freemium pour le B2C : Free (acquisition), Premium 9.99EUR/mois (conversion), Premium+ "
    "19.99EUR/mois (valeur premium). B2B en marque blanche : les partenaires (salles de sport, "
    "mutuelles, entreprises) paient une licence pour deployer leur version personnalisee a leurs "
    "clients/employes - ce qui justifie le tier Premium+ et la page Clients."
)

add_qa(doc,
    "Quelles sont les frictions dans l'adoption ?",
    "1) Inscription : un wizard de 5 etapes peut faire abandonner. J'ai ajoute la persistance "
    "localStorage pour reprendre ou on en etait. 2) Confidentialite : utilisateurs reticents a "
    "donner leur poids/taille. La FAQ et la note IMC adresseront cela. 3) Upgrade premium : il faut "
    "convaincre. Le PremiumGuard avec aperçu floute du contenu vise a creer l'envie."
)

# === 14.9 LIMITES ===
add_heading(doc, "14.9 Limites du projet et axes d'amelioration", 2)

add_qa(doc,
    "Quelles sont les principales limites de votre travail ?",
    "1) Pas de role admin implemente : la page Clients filtre par Premium+ mais aucun controle de "
    "role. 2) Pas de tests unitaires (uniquement E2E). 3) Pas d'i18n. 4) Pas de PWA web. 5) Mode "
    "demo qui peut masquer des bugs. 6) Pas de monitoring en prod."
)

add_qa(doc,
    "Le mode demo n'est-il pas un risque ?",
    "Si - il peut masquer un bug d'integration backend. Pour le mitiger : 1) Toujours tester avec "
    "backend en preprod avant deploiement. 2) Ajouter un bandeau visuel 'Mode demo' quand le "
    "fallback est actif. 3) Logger les erreurs pour ne pas qu'elles passent sous le radar. "
    "Aujourd'hui le bandeau est en place sur les pages critiques."
)

add_qa(doc,
    "Avez-vous teste sur de vrais utilisateurs ?",
    "Tests utilisateurs limites dans le cadre du projet etudiant. En conduite du changement, "
    "j'ai prevu : 1) Un panel beta de 10-20 utilisateurs avant lancement. 2) Tracking des "
    "evenements (inscription complete, abandon par etape) avec un outil comme Mixpanel ou "
    "Posthog pour identifier les frictions. 3) Iterations courtes apres lancement."
)

# === 14.10 PIEGES ===
add_heading(doc, "14.10 Questions pieges potentielles", 2)

add_qa(doc,
    "Pourquoi n'avez-vous pas utilise des Server Components partout pour le SEO ?",
    "Vrai point - certaines pages dashboard pourraient etre Server Components. Je les ai faites "
    "Client Components car elles ont besoin d'etat (loading, fallback demo, hooks usePremiumStatus). "
    "Pour la landing page (page.tsx), c'est en effet un Server Component, ce qui aide au SEO. "
    "Pour les pages dashboard, le SEO est moins critique car elles sont derriere un login."
)

add_qa(doc,
    "Quelle est la complexite algorithmique du calcul de l'IMC ?",
    "O(1) - c'est une simple division. La 'complexite' interessante est ailleurs : la classification "
    "en 4 categories est aussi O(1), mais l'interpretation pour un utilisateur reel demande plus "
    "que l'IMC (composition corporelle, ratio taille/hanches, etc.) - c'est pour ca que j'ai ajoute "
    "la note explicative."
)

add_qa(doc,
    "Et si le backend etait down en production ?",
    "1) Le mode demo permet a l'app de continuer a afficher quelque chose (utile pour un magasin "
    "de demo). 2) Un bandeau 'Mode demo' previent l'utilisateur que les donnees ne sont pas reelles. "
    "3) Sentry alerterait l'equipe en quelques minutes. 4) On afficherait une page de status "
    "(status.health-ai.com) communiquant l'incident. 5) Idealement, un load balancer ferait basculer "
    "vers une replique."
)

add_qa(doc,
    "Que se passe-t-il si l'utilisateur a JavaScript desactive ?",
    "Next.js fait du SSR : le HTML initial est servi cote serveur, donc la page s'affiche meme sans "
    "JS. Mais l'app n'est pas interactive (pas de clic, pas de fetch). Pour un dashboard, ca limite "
    "fort. Une amelioration serait d'identifier les pages qui peuvent fonctionner en HTML+CSS pur "
    "(landing, connexion) et de progressivement enrichir avec JS - principe de progressive enhancement."
)

add_qa(doc,
    "Comment garantir que Tailwind ne va pas exploser le CSS final ?",
    "Tailwind utilise du tree-shaking : le compilateur scanne les fichiers source et ne genere "
    "que les classes effectivement utilisees. Le CSS final fait typiquement 10-30 Ko gzippe. Si "
    "on utilisait des classes generes dynamiquement (template literals), il faudrait les whitelister "
    "dans tailwind.config car le scanner statique ne les detecterait pas."
)

add_qa(doc,
    "Comment debugger un probleme de performance sur une page ?",
    "1) Chrome DevTools Performance tab : enregistrer une session, identifier les long tasks. "
    "2) React DevTools Profiler : voir quels composants se re-rendent inutilement. 3) Lighthouse : "
    "auditer LCP, FID, CLS. 4) Bundle Analyzer (next-bundle-analyzer) pour reperer un import lourd. "
    "5) Network tab pour voir les requetes lentes ou bloquantes."
)

add_qa(doc,
    "Comment fonctionnerait votre application sur un mobile a connexion lente ?",
    "Next.js compresse les assets en production (gzip/brotli). Les images sont servies en WebP/AVIF "
    "via next/image avec lazy loading et placeholder. Le code splitting reduit le bundle initial. "
    "Le mode demo permet meme un fonctionnement degrade. Pour aller plus loin : Service Worker "
    "(PWA) pour cacher les assets statiques, et garder un mode offline minimal."
)

add_qa(doc,
    "Vous avez Premium+ mais aucun controle d'admin reel. C'est un probleme de securite ?",
    "Pas dans la version actuelle car la page Clients ne montre aucune donnee personnelle (juste "
    "des stats agregees anonymes). Mais oui, c'est un trou si demain on ajoute la liste nominative. "
    "Il faut alors ajouter un role 'admin' verifie cote backend (claim dans le JWT) et un check "
    "frontend correspondant. C'est explicitement marque dans les axes d'amelioration."
)

add_qa(doc,
    "Pourquoi vos branches features ont des noms en francais ?",
    "Decision pragmatique : le projet est un projet etudiant francais, lu par un jury francais. "
    "L'uniformite linguistique des branches/commits/docs facilite la lecture. En entreprise "
    "internationale, je passerais a l'anglais pour faciliter la collaboration multi-equipes."
)

add_qa(doc,
    "Vous parlez d'IA - mais quelle IA y a-t-il dans le frontend ?",
    "Le frontend ne fait pas d'IA lui-meme. Il consomme deux services IA : 1) ia-python (PyTorch "
    "ResNet18/MobileNet) pour l'analyse d'images de repas - on envoie une photo, on recoit la "
    "classification. 2) engine-go pour les recommandations (meal plans, workouts), qui utilise "
    "pgvector pour de la recherche par similarite vectorielle. Le frontend orchestre l'experience "
    "utilisateur autour de ces capacites IA."
)

add_qa(doc,
    "Si le jury vous demande une demo en live, vous etes pret ?",
    "Oui : 1) `npm run dev` lance le frontend sur localhost:3000. 2) Mode demo si le backend n'est "
    "pas dispo - aucune dependance externe necessaire. 3) Je peux montrer : landing -> wizard -> "
    "dashboard -> nutrition -> meal plan (PremiumGuard) -> workouts (PremiumGuard) -> clients "
    "(filtre Premium+). 4) Je peux montrer le code derriere chaque page (VS Code preconfigure)."
)

doc.add_page_break()

# === ANNEXES ===
add_heading(doc, "Annexes", 1)

add_heading(doc, "A. Resume executif (a memoriser)", 2)
add_para(doc,
    "Health AI est une plateforme de coaching sante avec analyse IA, declinee en frontend Next.js 16, "
    "mobile Flutter, BFF Hono, moteur Go, et service IA Python. Mon volet (frontend) couvre 10+ "
    "pages, un wizard d'inscription persistant, un systeme de tiers (Free / Premium / Premium+) "
    "avec guards reutilisables, l'integration Better Auth, des graphiques Recharts, l'accessibilite "
    "WCAG/RGAA AA testee automatiquement, des tests E2E Playwright, et une documentation complete "
    "(11 etapes + 6 livrables). Stack : Next.js 16, TypeScript, Tailwind 4, Shadcn UI, Zustand, "
    "Better Auth, Playwright, axe-core."
)

add_heading(doc, "B. Chiffres cles", 2)
for c in [
    "10+ pages developpees.",
    "5 etapes wizard d'inscription.",
    "3 tiers d'abonnement.",
    "4 themes B2B.",
    "14 composants Shadcn UI integres.",
    "23 screenshots dans docs/screenshots/.",
    "11 documents par etape (docs/01 a docs/11).",
    "6 documents de livraison (benchmark, tech, design, a11y, conduite changement, tests).",
    "46+ tests E2E Playwright.",
    "6 pages auditees automatiquement par axe-core.",
    "WCAG 2.1 AA / RGAA AA respecte.",
    "87+ commits sur la branche main.",
]:
    add_bullet(doc, c)

add_heading(doc, "C. Phrases d'accroche pour l'oral", 2)
add_para(doc, "Ouverture :", bold=True)
add_para(doc,
    "\"Health AI est une plateforme de coaching sante propulsee par l'IA. Mon travail a porte sur "
    "le frontend web : 10 pages, un wizard d'inscription, un systeme de tiers freemium / B2B, "
    "le tout pense pour l'accessibilite et teste automatiquement.\""
)
add_para(doc, "Cloture :", bold=True)
add_para(doc,
    "\"En 12 semaines, j'ai livre un frontend complet, accessible, teste et documente. Les principales "
    "limites - role admin, monitoring, i18n - sont identifiees et constituent les prochains chantiers. "
    "Je suis pret a vos questions.\""
)

# === SAUVEGARDE ===
output_path = r"C:\Developper\Ecole\Health-AI-project\preparation_oral_frontend.docx"
doc.save(output_path)
print(f"Document genere : {output_path}")
