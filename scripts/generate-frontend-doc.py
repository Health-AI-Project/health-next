from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

p = doc.add_paragraph("Documentation technique")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Frontend Health-Next")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Application web moderne, responsive et accessible pour HealthAI Coach")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Pierre Parain, Mehdi, Alexandre, Kylian")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("Introduction", level=1)
doc.add_paragraph(
    "Le frontend de HealthAI Coach est une application web construite avec Next.js 16, React 19 et TypeScript. "
    "Elle sert d'interface principale pour les utilisateurs finaux, les partenaires B2B et les administrateurs. "
    "L'objectif est de fournir une experience fluide et accessible pour consulter les recommandations IA en nutrition "
    "et activite physique, tout en respectant les normes d'accessibilite WCAG/RGAA niveau AA."
)
doc.add_paragraph(
    "L'application suit une architecture modulaire basee sur le systeme App Router de Next.js. "
    "Chaque page du dashboard est un composant client autonome qui fetch ses donnees depuis l'API backend "
    "et affiche des donnees de demonstration en fallback quand le serveur n'est pas disponible. "
    "Cela permet de tester et presenter l'application sans aucune dependance externe."
)
doc.add_paragraph(
    "L'interface utilise Shadcn UI comme bibliotheque de composants, construite sur Radix UI pour l'accessibilite "
    "et Tailwind CSS 4 pour le styling responsive. Un systeme de guard gere les restrictions d'acces selon trois "
    "tiers d'abonnement (Freemium, Premium, Premium+), conformement au business model du cahier des charges."
)

doc.add_heading("Architecture du projet", level=1)
doc.add_paragraph(
    "L'application est organisee autour du systeme de fichiers de Next.js App Router. Chaque dossier dans app/ "
    "correspond a une route. Le layout du dashboard (app/dashboard/layout.tsx) encapsule toutes les pages privees "
    "dans un DashboardLayout avec sidebar, theme provider et toaster de notifications."
)
doc.add_paragraph(
    "Les composants sont repartis en plusieurs categories. Le dossier components/ui/ contient les 14 composants Shadcn "
    "(button, card, input, form, select, checkbox, label, dialog, tabs, table, badge, progress, skeleton, toaster). "
    "Le dossier components/charts/ contient les trois graphiques Recharts (poids, calories, macronutriments). "
    "Le dossier components/nutrition/ contient les composants metier lies a l'analyse nutritionnelle. "
    "Le dossier components/premium/ contient le systeme de guard (PremiumGuard et UpgradeDialog). "
    "Le dossier components/wizard/ contient les etapes du formulaire d'inscription."
)
doc.add_paragraph(
    "La couche donnees repose sur un utilitaire apiFetch() dans lib/api.ts qui centralise tous les appels HTTP "
    "avec gestion automatique des headers JSON, des credentials et des erreurs. Les schemas de validation Zod "
    "sont dans lib/schemas/. L'etat global du wizard est gere par un store Zustand dans lib/stores/. "
    "Le systeme de themes B2B est dans lib/themes/ avec quatre themes preconfigures."
)

doc.add_heading("Pages et routes", level=1)

doc.add_heading("Landing page (/)", level=2)
doc.add_paragraph(
    "La page d'accueil est publique et presente l'offre HealthAI Coach. Elle affiche un hero avec le titre et la proposition "
    "de valeur, puis trois cards de pricing cote a cote : Freemium (0 euros/mois), Premium (9,99 euros/mois) et Premium+ (19,99 euros/mois). "
    "Chaque card liste ses fonctionnalites avec des icones check/cross et un bouton CTA qui redirige vers l'inscription. "
    "La card Premium porte un badge \"Populaire\" et la card Premium+ utilise la variante bouton gradient. "
    "Les trois cards ont un effet de hover avec scale et shadow progressive."
)

doc.add_heading("Inscription (/inscription)", level=2)
doc.add_paragraph(
    "Le wizard d'inscription est un formulaire en 6 etapes : age, poids et taille, objectifs de sante, allergies alimentaires, "
    "recapitulatif et creation de compte. Chaque etape est un composant autonome avec validation Zod. "
    "La navigation entre les etapes est geree par le composant WizardNavigation (boutons Precedent/Suivant) "
    "et l'etat est persiste dans un store Zustand. La barre de progression indique l'avancement. "
    "Le step recapitulatif affiche les informations sous forme de SummaryCards et le step final "
    "envoie les donnees a l'API d'authentification Better Auth."
)

doc.add_heading("Dashboard principal (/dashboard)", level=2)
doc.add_paragraph(
    "Le dashboard est la page d'accueil de l'espace connecte. Il affiche un message de bienvenue personnalise "
    "avec l'email de l'utilisateur, quatre cards de statistiques cles (poids actuel, calories du jour, proteines, activite), "
    "une card IMC calculee a partir du poids et de la taille, puis deux graphiques interactifs "
    "(evolution du poids sur 30 jours et calories journalieres de la semaine). "
    "La card IMC affiche la valeur avec une interpretation coloree en quatre niveaux (insuffisance ponderale en bleu, "
    "poids normal en vert, surpoids en orange, obesite en rouge), une barre de progression et une note rappelant "
    "que l'IMC ne distingue pas masse musculaire et masse grasse."
)

doc.add_heading("Analytics (/dashboard/analytics)", level=2)
doc.add_paragraph(
    "La page analytics offre une vue detaillee des tendances de l'utilisateur. Elle presente quatre cards de resume "
    "(calories moyennes, proteines moyennes, seances effectuees, poids actuel) puis trois onglets : "
    "Vue d'ensemble (trois charts), Nutrition (calories + macros) et Poids (evolution seule). "
    "Le graphique des macronutriments est un PieChart donut montrant la repartition proteines/glucides/lipides. "
    "Ce chart est protege par un PremiumGuard : les utilisateurs gratuits voient le contenu floute "
    "avec un bouton pour debloquer la fonctionnalite."
)

doc.add_heading("Nutrition (/dashboard/nutrition)", level=2)
doc.add_paragraph(
    "La page nutrition est le point d'entree pour l'analyse de repas par IA. L'utilisateur uploade une photo "
    "de son repas via une zone de drag-and-drop (react-dropzone) qui accepte les formats PNG, JPG et WEBP jusqu'a 10 Mo. "
    "L'image est envoyee a l'API backend en FormData POST. Le resultat s'affiche dans un tableau editable "
    "listant les aliments detectes avec leurs calories, proteines, glucides et lipides, ainsi qu'un total."
)
doc.add_paragraph(
    "Apres l'analyse, une section Analyse nutritionnelle apparait automatiquement. Elle compare les totaux du repas "
    "aux objectifs journaliers divises par trois (un repas represente un tiers de la journee). "
    "Chaque nutriment recoit un statut : deficit (moins de 80% de l'objectif, badge orange), "
    "equilibre (entre 80 et 120%, badge vert) ou exces (plus de 120%, badge rouge). "
    "Pour chaque desequilibre, une suggestion personnalisee est affichee, par exemple "
    "\"Ajoutez une source de proteines : poulet, poisson, oeufs, legumineuses ou tofu\" en cas de deficit proteique."
)

doc.add_heading("Journal alimentaire (/dashboard/nutrition/history)", level=2)
doc.add_paragraph(
    "Le journal alimentaire affiche l'historique de tous les repas analyses. Trois cards de resume "
    "indiquent le nombre de repas enregistres, les calories totales et la moyenne par repas. "
    "Le tableau liste chaque repas avec sa date, son type (badge Dejeuner, Diner, etc.), "
    "les aliments detectes et les macronutriments. L'utilisateur peut filtrer par periode "
    "(Aujourd'hui, 7 jours, 30 jours, Tout) via des onglets et trier par date via un bouton bascule."
)

doc.add_heading("Plans de repas (/dashboard/nutrition/meal-plan)", level=2)
doc.add_paragraph(
    "La page des plans de repas affiche un plan nutritionnel hebdomadaire genere par l'IA. "
    "Cinq jours sont presentes sous forme d'onglets (Lundi a Vendredi). Chaque jour contient trois a quatre repas "
    "(petit-dejeuner, dejeuner, diner, collation optionnelle) affiches en cards avec le nom du plat, "
    "les calories en badge, les macronutriments et la liste des ingredients en badges secondaires. "
    "Le total calorique du jour est affiche en haut. Cette page est entierement protegee par un PremiumGuard : "
    "les utilisateurs gratuits voient le contenu floute avec un dialog d'upgrade vers Premium."
)

doc.add_heading("Entrainement (/dashboard/workouts)", level=2)
doc.add_paragraph(
    "La page entrainement affiche un programme d'entrainement hebdomadaire personnalise par l'IA. "
    "Trois cards de statistiques en haut indiquent le nombre de seances par semaine, les calories brulees "
    "et le nombre total d'exercices. Le programme est organise en cinq jours sous forme d'onglets, "
    "chacun avec un focus musculaire (haut du corps push, bas du corps, repos actif, haut du corps pull, full body HIIT), "
    "une duree, les calories estimees et un badge de difficulte (Debutant, Intermediaire, Avance). "
    "Chaque exercice est affiche dans une card individuelle avec un numero d'ordre, le nom de l'exercice, "
    "le nombre de series et repetitions, le temps de repos et le muscle cible en badge. "
    "Le programme est protege par un PremiumGuard : les statistiques restent visibles pour tous, "
    "mais les exercices sont floutes pour les utilisateurs gratuits."
)

doc.add_heading("Parametres (/dashboard/settings)", level=2)
doc.add_paragraph(
    "La page parametres est organisee en trois onglets avec icones : Profil, Objectifs et Abonnement. "
    "L'onglet Profil permet de modifier l'age, le poids et la taille (l'email est affiche en lecture seule). "
    "L'onglet Objectifs permet de cocher/decocher les six objectifs de sante et les huit allergies alimentaires, "
    "avec la logique speciale \"Aucune allergie\" qui decoche toutes les autres. "
    "L'onglet Abonnement affiche le tier actuel avec un badge (Freemium, Premium ou Premium+) "
    "et propose les upgrades pertinents : un utilisateur gratuit voit les deux offres Premium et Premium+ "
    "cote a cote, un utilisateur Premium voit l'offre Premium+ uniquement, et un utilisateur Premium+ "
    "voit un message de confirmation."
)

doc.add_heading("Clients B2B (/dashboard/clients)", level=2)
doc.add_paragraph(
    "La page clients est reservee aux comptes B2B et Premium+. Les utilisateurs gratuits et Premium "
    "voient un ecran d'acces reserve avec un bouton de redirection vers les offres. "
    "Les utilisateurs Premium+ voient quatre cards de statistiques agregees : total clients, "
    "clients actifs sur les sept derniers jours, taux de conversion Premium et calories moyennes. "
    "Les donnees personnelles des clients (noms, emails) ne sont pas affichees : "
    "cette fonctionnalite est preparee pour etre branchee quand le backend exposera un role administrateur."
)

doc.add_heading("Systeme de guard (Freemium/Premium/Premium+)", level=1)
doc.add_paragraph(
    "Le systeme de guard repose sur trois composants. Le hook usePremiumStatus() dans lib/hooks/ "
    "appelle l'API /api/home au montage du composant et extrait le tier de l'utilisateur. "
    "Il supporte deux champs de l'API : subscription_tier (nouveau, prioritaire) et is_premium (retrocompatible). "
    "Le hook retourne le tier actuel, un boolean isLoading, et trois helpers : isPremium, isPremiumPlus et canAccess(requiredTier). "
    "La comparaison se fait via un niveau numerique : free=0, premium=1, premium_plus=2."
)
doc.add_paragraph(
    "Le composant PremiumGuard dans components/premium/ accepte un children (contenu a proteger), "
    "un feature (nom de la fonctionnalite pour le dialog) et un requiredTier optionnel (defaut premium). "
    "Pendant le chargement, il affiche un Skeleton. Si l'utilisateur a le tier suffisant, "
    "il affiche le children normalement. Sinon, il affiche le contenu floute avec opacite reduite "
    "et un overlay centre contenant le bouton d'upgrade."
)
doc.add_paragraph(
    "Le composant UpgradeDialog dans components/premium/ est un Dialog Radix qui s'adapte au tier requis. "
    "Si la fonctionnalite est Premium, il affiche le prix de 9,99 euros, l'icone couronne et les trois features Premium. "
    "Si la fonctionnalite est Premium+, il affiche le prix de 19,99 euros, l'icone diamant et les trois features Premium+. "
    "Le message s'adapte aussi au tier actuel de l'utilisateur."
)

doc.add_heading("Marque blanche B2B", level=1)
doc.add_paragraph(
    "Le systeme de marque blanche permet aux partenaires B2B de personnaliser l'apparence de la plateforme. "
    "Quatre themes sont preconfigures dans lib/themes/company-themes.ts : HealthNext (vert), Gym Club (violet), "
    "FitLife Pro (orange) et Wellness Center (cyan). Chaque theme definit un nom, une icone, et des couleurs HSL "
    "pour le primary, l'accent et les trois couleurs de graphiques."
)
doc.add_paragraph(
    "Le DynamicThemeProvider applique les couleurs du theme selectionne en modifiant les variables CSS "
    "sur l'element racine du document. Il gere egalement le mode de couleur (clair, sombre, systeme) "
    "avec detection de la preference systeme via media query. Le theme et le mode sont persistes dans le localStorage."
)

doc.add_heading("Composants UI (Shadcn)", level=1)
doc.add_paragraph(
    "L'application utilise 14 composants Shadcn UI dans le dossier components/ui/. Chaque composant est base "
    "sur une primitive Radix UI pour l'accessibilite et style avec Tailwind CSS via class-variance-authority (CVA). "
    "Les composants sont : button (avec variante premium en gradient), card, checkbox, dialog, form, input, label, "
    "progress, select, skeleton, table, tabs, badge (4 variantes) et toaster (wrapper Sonner). "
    "Tous les composants interactifs utilisent focus-visible:ring-2 pour l'indication de focus clavier."
)

doc.add_heading("Visualisations interactives", level=1)
doc.add_paragraph(
    "Trois graphiques interactifs sont implementes avec Recharts. Le WeightEvolutionChart est un LineChart "
    "a deux lignes (poids actuel et objectif en pointilles) avec des tooltips au survol. "
    "Le CaloriesChart est un BarChart avec barres colorees selon l'objectif et une ligne de reference a 2000 kcal. "
    "Le MacrosChart est un PieChart en donut montrant la repartition des macronutriments. "
    "Tous les graphiques s'adaptent au theme B2B selectionne grace au hook useChartColors()."
)

doc.add_heading("Integration API et fallback", level=1)
doc.add_paragraph(
    "Toutes les pages du dashboard suivent le meme pattern d'integration API. Au montage du composant, "
    "un useEffect appelle apiFetch() qui envoie une requete HTTP avec credentials inclus. "
    "En cas de succes, les donnees de l'API sont affichees. En cas d'erreur (backend absent, timeout, erreur serveur), "
    "le catch charge des donnees de demonstration definies en constante dans le fichier de la page. "
    "Pendant le chargement, des Skeletons s'affichent a la place du contenu."
)

doc.add_heading("Tests automatises et accessibilite", level=1)
doc.add_paragraph(
    "Les tests sont executes avec Playwright et couvrent deux axes. Les tests E2E dans e2e/user-journey.spec.ts "
    "verifient le parcours utilisateur critique : navigation landing vers inscription, deroulement du wizard, "
    "validation des erreurs de formulaire et navigation dans le dashboard."
)
doc.add_paragraph(
    "Les tests d'accessibilite dans e2e/accessibility.spec.ts utilisent axe-core pour scanner six pages "
    "contre les criteres WCAG 2.1 AA. Les resultats sont sauvegardes en JSON dans le dossier a11y-report/ "
    "et un script Node.js genere un rapport HTML presentable avec un resume global, "
    "le detail par page, les violations par severite et un tableau de correspondance WCAG vers RGAA. "
    "La commande npm run test:a11y:report lance les tests et genere le rapport en une seule etape."
)

doc.add_heading("Installation et prise en main", level=1)

doc.add_heading("Prerequis", level=2)
doc.add_paragraph("Le projet requiert Node.js 20 ou superieur et les navigateurs Chromium pour les tests Playwright.")

doc.add_heading("Installation", level=2)
doc.add_paragraph("cd health-next")
doc.add_paragraph("npm install")
doc.add_paragraph("npx playwright install --with-deps chromium")

doc.add_heading("Commandes disponibles", level=2)
doc.add_paragraph("npm run dev : serveur de developpement avec hot reload")
doc.add_paragraph("npm run build : compilation de production")
doc.add_paragraph("npm run lint : verification ESLint")
doc.add_paragraph("npm run test : tous les tests Playwright")
doc.add_paragraph("npm run test:e2e : tests parcours utilisateur")
doc.add_paragraph("npm run test:a11y : tests accessibilite")
doc.add_paragraph("npm run test:a11y:report : tests accessibilite + generation rapport HTML")

doc.add_heading("Limites et points d'attention", level=1)
doc.add_paragraph(
    "Les donnees de demonstration sont statiques et codees en dur dans chaque page. Elles ne sont pas synchronisees "
    "entre les pages. Quand le backend sera connecte, ces constantes ne seront plus utilisees."
)
doc.add_paragraph(
    "Le systeme de guard repose sur un appel API a chaque montage de composant PremiumGuard. "
    "Une optimisation future pourrait centraliser le tier dans un contexte React partage."
)
doc.add_paragraph(
    "Les donnees personnelles de la page clients ne sont pas encore affichees pour les administrateurs "
    "car le backend n'expose pas encore de role admin."
)
doc.add_paragraph(
    "L'estimation des desequilibres nutritionnels utilise des seuils fixes (2000 kcal/jour, 60g proteines, etc.) "
    "qui ne tiennent pas compte du profil specifique de l'utilisateur. Une amelioration future pourrait "
    "adapter ces seuils en fonction de l'age, du poids, de la taille et de l'objectif de l'utilisateur."
)

doc.save(r"C:\Developper\Ecole\Health-AI-project\documentation_frontend.docx")
print("Document cree avec succes")
