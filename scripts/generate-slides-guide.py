"""
Genere un document guide pour le diaporama de la partie FRONTEND
de l'oral MSPR. Indique pour chaque slide : titre, contenu, visuel,
duree, et notes de presentateur.

Sortie : .docx + .pdf dans C:\\Developper\\Ecole\\Health-AI-project\\
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

OUT_DIR = Path(r"C:\Developper\Ecole\Health-AI-project")
DOCX = OUT_DIR / "guide_diapo_frontend.docx"
PDF = OUT_DIR / "guide_diapo_frontend.pdf"

PRIMARY = RGBColor(0x1F, 0x3A, 0x8A)
ACCENT = RGBColor(0x0E, 0x76, 0x90)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_base_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = PRIMARY


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = PRIMARY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def add_field(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label} : ")
    r1.bold = True
    r1.font.color.rgb = ACCENT
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def add_bullets(doc, label, items):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label} :")
    r1.bold = True
    r1.font.color.rgb = ACCENT
    r1.font.size = Pt(11)
    for it in items:
        bp = doc.add_paragraph(it, style="List Bullet")
        bp.paragraph_format.space_after = Pt(0)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)


def add_slide(doc, num, total, titre, duree, contenu, visuel, notes):
    add_h2(doc, f"Slide {num}/{total} - {titre}")
    add_field(doc, "Duree", duree)
    add_bullets(doc, "Contenu a afficher", contenu)
    add_field(doc, "Visuel suggere", visuel)
    add_field(doc, "Notes orateur", notes)


def main():
    doc = Document()
    set_base_style(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_title(doc, "Guide du diaporama - Frontend Health AI")
    add_subtitle(doc, "Oral MSPR TPRE502 Bac+3 - 5 minutes - 8 slides")
    doc.add_paragraph()

    add_para(
        doc,
        "Ce document decrit les slides a creer pour la presentation de la partie frontend. "
        "Le diaporama suit le script oral et tient en 8 slides pour respecter les 5 minutes. "
        "Regle d'or : peu de texte par slide, une idee par slide, des visuels qui parlent. "
        "Le jury doit ECOUTER, pas LIRE. Chaque slide est minutee."
    )

    total = 8

    add_slide(
        doc, 1, total,
        "Slide de couverture",
        "10 secondes",
        [
            "Titre : Health AI Platform - Frontend Web",
            "Sous-titre : MSPR TPRE502 - Bac+3",
            "Votre nom et la date de l'oral",
            "Logo / nom de l'ecole",
        ],
        "Logo du projet ou capture d'ecran du dashboard en arriere-plan flouté.",
        "Saluer le jury, donner votre nom, annoncer le sujet. Pas plus.",
    )

    add_slide(
        doc, 2, total,
        "Contexte et objectifs",
        "30 secondes",
        [
            "Health AI : application de coaching sante avec IA",
            "Mon perimetre : interface web (Next.js)",
            "Objectif : un dashboard accessible, responsive, avec 3 niveaux d'abonnement",
            "Public cible : utilisateurs grand public + administrateurs",
        ],
        "Schema simple : utilisateur -> frontend web -> BFF -> services backend. Mettre en avant la zone frontend en couleur, le reste grise.",
        "Poser le decor en 30 secondes. Insister sur le fait que vous parlez UNIQUEMENT de la couche web.",
    )

    add_slide(
        doc, 3, total,
        "Stack technique",
        "45 secondes",
        [
            "Next.js 16 (App Router) + React 19 + TypeScript strict",
            "Tailwind CSS 4 + Shadcn UI / Radix UI",
            "Zustand (etat) + React Hook Form + Zod (formulaires)",
            "Better Auth (authentification) + Recharts (graphiques)",
        ],
        "Logos en grille 2x4 ou 4x2 (Next.js, React, TypeScript, Tailwind, Shadcn, Zustand, Zod, Recharts). Pas de phrases, juste les logos.",
        "Justifier les choix : typage fort, accessibilite native via Radix, productivite. Ne PAS lire la liste, paraphraser.",
    )

    add_slide(
        doc, 4, total,
        "Architecture frontend",
        "1 minute",
        [
            "Le frontend ne parle qu'au BFF Hono via REST",
            "Aucun acces direct aux bases de donnees ni au service Go",
            "Pattern resilient : appel API -> fallback donnees demo en cas d'echec",
            "App Router avec route groups : (auth) public, dashboard protege",
        ],
        "Diagramme : [Frontend Next.js] --REST--> [BFF Hono] --gRPC--> [Engine Go] --> [PostgreSQL]. Encadrer la partie frontend.",
        "Insister sur la separation des responsabilites et le pattern de fallback qui permet la demo sans backend.",
    )

    add_slide(
        doc, 5, total,
        "Pages et fonctionnalites cles",
        "1 minute",
        [
            "Landing + auth multi-etapes avec onboarding",
            "Dashboard : resume nutritionnel, IMC, suggestions repas",
            "Nutrition : analyse photo, historique, meal plan hebdomadaire",
            "Workouts, Analytics (graphiques), Clients (admin), Parametres",
        ],
        "Mosaique de 4 a 6 captures d'ecran reelles : dashboard, meal plan, analytics, page clients admin. Petit titre sous chaque capture.",
        "Ne pas detailler chaque page. Citer les plus impressionnantes (analytics avec Recharts, page clients admin).",
    )

    add_slide(
        doc, 6, total,
        "Roles et abonnements",
        "45 secondes",
        [
            "Tier d'abonnement : free / premium / premium_plus",
            "Role technique : user / admin (independant du tier)",
            "PremiumGuard verrouille les pages selon le tier",
            "Hook useUserRole + mode demo via localStorage pour la presentation",
        ],
        "Tableau 2 colonnes : a gauche les tiers (free, premium, premium+), a droite les roles (user, admin). Fleche montrant que ce sont 2 axes orthogonaux.",
        "Bien expliquer la difference entre 'ce que je paie' (tier) et 'ce que je peux faire techniquement' (role).",
    )

    add_slide(
        doc, 7, total,
        "Accessibilite et qualite",
        "45 secondes",
        [
            "Conformite WCAG 2.1 AA / RGAA AA",
            "Labels ARIA, navigation clavier, contrastes verifies",
            "Tests Playwright (E2E) + axe-core (a11y automatise)",
            "CI GitHub Actions sur chaque pull request",
        ],
        "3 badges/icones : badge WCAG AA, logo Playwright, logo GitHub Actions avec coche verte. Ou un screenshot du rapport axe-core.",
        "L'accessibilite est differenciante a l'oral, beaucoup d'eleves zappent. Insister 5 secondes dessus.",
    )

    add_slide(
        doc, 8, total,
        "Conclusion + demo",
        "30 secondes",
        [
            "Frontend moderne, typee, accessible, testee",
            "Mode demo pret pour presentation live",
            "Prochaines etapes : i18n, PWA offline, monitoring",
            "Merci - questions ?",
        ],
        "Capture du dashboard final en grand, semi-transparente avec le mot 'Demo' ou 'Questions ?' en surimpression.",
        "Enchainer sur la demo live si prevue. Sinon, ouvrir aux questions avec le sourire.",
    )

    add_h2(doc, "Conseils generaux de mise en forme")
    for tip in [
        "Charte graphique : 2 couleurs principales + 1 accent. Reprendre les couleurs du dashboard pour la coherence.",
        "Typographie : un seul jeu de polices (sans-serif moderne, type Inter ou Calibri). Titres 32-40 pt, corps 18-22 pt.",
        "Densite : maximum 5 a 6 lignes de texte par slide, jamais de paragraphe.",
        "Captures d'ecran : prendre des vraies captures du projet, pas de mockups generiques. Cropper proprement.",
        "Animations : aucune, ou tres legeres (fondu). Pas de transitions agressives qui distraient le jury.",
        "Numero de slide visible en bas a droite + nom du projet en bas a gauche pour aider le jury a se reperer.",
        "Repetez l'oral 2 fois chronometre AVEC le diapo, pas separement. Vous verrez naturellement quelles slides ralentissent.",
    ]:
        doc.add_paragraph(tip, style="List Bullet")

    add_h2(doc, "Outils suggeres")
    for tool in [
        "PowerPoint, Keynote, ou Google Slides : standards, fiables.",
        "Canva : si vous voulez un rendu plus design rapidement, plein de templates Bac+3 propres.",
        "Pitch.com ou Slidev : alternatives modernes, Slidev permet de coder les slides en markdown si vous etes a l'aise.",
    ]:
        doc.add_paragraph(tool, style="List Bullet")

    doc.save(str(DOCX))
    print(f"docx genere : {DOCX}")
    convert(str(DOCX), str(PDF))
    print(f"pdf genere : {PDF}")


if __name__ == "__main__":
    main()
