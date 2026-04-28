"""
Genere un script d'oral MSPR (Bac+3) de 5 minutes maxi sur la partie FRONTEND
du projet Health AI Platform. Sortie : .docx + .pdf

Cible : ~750 mots (debit ~150 mots/min).
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

OUT_DIR = Path(r"C:\Developper\Ecole\Health-AI-project")
DOCX = OUT_DIR / "oral_frontend_5min.docx"
PDF = OUT_DIR / "oral_frontend_5min.pdf"


def set_base_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_section(doc, label, body):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)
    para = doc.add_paragraph(body)
    para.paragraph_format.space_after = Pt(8)
    for r in para.runs:
        r.font.size = Pt(11)


def main():
    doc = Document()
    set_base_style(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_title(doc, "Oral MSPR TPRE502 - Frontend Health AI")
    add_subtitle(doc, "Script de presentation - 5 minutes - Bac+3")
    doc.add_paragraph()

    add_section(
        doc,
        "1. Introduction (30 s)",
        "Bonjour, je vais vous presenter la partie frontend web du projet Health AI Platform, "
        "une application de coaching sante alimentee par de l'intelligence artificielle. "
        "Mon role a ete de concevoir et developper l'interface web complete, accessible et "
        "responsive, qui consomme les services backend via une API REST. L'objectif : offrir "
        "aux utilisateurs un dashboard clair pour suivre leur nutrition, leurs entrainements "
        "et leur progression, avec une experience adaptee a trois niveaux d'abonnement.",
    )

    add_section(
        doc,
        "2. Stack technique (45 s)",
        "J'ai choisi Next.js 16 avec l'App Router, TypeScript strict et React 19. "
        "Pour le style, Tailwind CSS 4 associe a la bibliotheque de composants Shadcn UI, "
        "construite sur Radix UI, ce qui garantit une accessibilite native sur les composants "
        "interactifs. La gestion d'etat se fait via Zustand pour les preferences utilisateur, "
        "avec persistance localStorage. Les formulaires utilisent React Hook Form et Zod pour "
        "la validation typee. L'authentification s'appuie sur Better Auth, et les graphiques "
        "sur Recharts. Ce choix de stack moderne, fortement typee, m'a permis de garantir la "
        "robustesse du code et une excellente productivite.",
    )

    add_section(
        doc,
        "3. Architecture et fonctionnalites (1 min 30)",
        "Le frontend communique uniquement avec le BFF Hono via REST. Il ne connait jamais "
        "directement le service Go ni la base de donnees, ce qui respecte le principe de "
        "separation des responsabilites dans une architecture microservices. "
        "J'ai construit dix pages principales : la landing page, l'authentification (connexion "
        "et inscription multi-etapes avec onboarding), le dashboard avec resume nutritionnel "
        "et IMC, la nutrition avec analyse photo et historique, le meal plan hebdomadaire, "
        "les workouts, les analytics avec graphiques d'evolution, la gestion des clients pour "
        "les administrateurs, et les parametres. "
        "Chaque page suit un pattern resilient : on essaie l'appel API, et en cas d'echec on "
        "bascule sur des donnees de demonstration. Cela permet de presenter l'application "
        "meme sans backend disponible, ce qui a ete crucial pour la phase de developpement et "
        "le sera aussi pendant cette demonstration.",
    )

    add_section(
        doc,
        "4. Roles, abonnements et acces (45 s)",
        "L'application gere deux dimensions distinctes : le tier d'abonnement - free, premium "
        "ou premium plus - qui ouvre l'acces aux fonctionnalites payantes comme les workouts "
        "ou les analytics, et le role technique - utilisateur ou administrateur - qui permet "
        "de gerer les autres utilisateurs. Un composant PremiumGuard verrouille les pages "
        "selon le tier, et un hook useUserRole detecte le role admin. Pour la demonstration "
        "sans backend, j'ai prevu un mode demo qui s'active via localStorage, ce qui me "
        "permet de basculer entre les vues en direct devant vous.",
    )

    add_section(
        doc,
        "5. Accessibilite et tests (45 s)",
        "L'accessibilite a ete une priorite : conformite RGAA AA et WCAG 2.1 AA. Tous les "
        "composants interactifs ont des labels ARIA, la navigation au clavier est complete, "
        "les contrastes respectent les ratios requis et les formulaires affichent des erreurs "
        "lisibles par les lecteurs d'ecran. "
        "J'ai mis en place des tests Playwright pour les parcours end-to-end et des tests "
        "axe-core automatises pour l'accessibilite. Une integration continue GitHub Actions "
        "execute le build et les tests sur chaque pull request, ce qui garantit la non "
        "regression.",
    )

    add_section(
        doc,
        "6. Conclusion (30 s)",
        "En resume, le frontend Health AI est une application web moderne, typee, accessible "
        "et testee, construite avec une stack robuste et une architecture claire. Le mode "
        "demo me permet de la presenter de maniere fluide. Je suis pret a vous faire une "
        "demonstration en direct et a repondre a vos questions. Merci.",
    )

    doc.save(str(DOCX))
    print(f"docx genere : {DOCX}")

    convert(str(DOCX), str(PDF))
    print(f"pdf genere : {PDF}")


if __name__ == "__main__":
    main()
